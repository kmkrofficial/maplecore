#!/usr/bin/env python3
"""
Verify Document Parsing
=======================
Tests MAPLE's ability to ingest rich PDF, DOCX, and HTML files.
"""

import os
from pathlib import Path
from maplecore import Maple

def create_dummy_files():
    """Create a basic .pdf, .docx, and .html file for testing"""
    files = []
    
    # 1. Provide an HTML
    html_path = "test_doc.html"
    with open(html_path, "w", encoding="utf-8") as f:
        f.write("<html><body><h1>MAPLE Parsing</h1><p>This is a secret HTML document about MAPLE.</p></body></html>")
    files.append(html_path)
    
    # 2. Provide a DOCX
    import docx
    doc_path = "test_doc.docx"
    doc = docx.Document()
    doc.add_heading('MAPLE Test Document', 0)
    doc.add_paragraph('This is a secret DOCX document about MAPLE.')
    doc.save(doc_path)
    files.append(doc_path)
    
    # 3. Provide a PDF
    import fitz # PyMuPDF
    pdf_path = "test_doc.pdf"
    pdf = fitz.open()
    page = pdf.new_page()
    page.insert_text((50, 50), "This is a secret PDF document about MAPLE.", fontsize=11)
    pdf.save(pdf_path)
    pdf.close()
    files.append(pdf_path)
    
    return files

def cleanup_files(files):
    for f in files:
        if os.path.exists(f):
            os.remove(f)

def run_test():
    print("Generating dummy files...")
    files = create_dummy_files()
    
    try:
        # Initialize client with CPU for lightweight test
        client = Maple(embedding_model="sentence-transformers/all-MiniLM-L6-v2", device="cpu")
        
        for file_path in files:
            print(f"\n--- Testing {Path(file_path).suffix.upper()} ---")
            
            # Index File
            print(f"Indexing {file_path}...")
            index = client.index_file(file_path, chunk_size=50)
            
            print(f"Index Blocks: {index.num_blocks}")
            
            # Verification
            found = False
            for bid in range(index.num_blocks):
                block = client.get_block(bid)
                if "secret" in block.text.lower():
                    found = True
                    print(f"Match found in block {bid}: {block.text}")
                    break
                    
            if found:
                print(f"{Path(file_path).suffix.upper()} Passed!")
            else:
                print(f"Failed: Target text not found in {file_path}")
                assert False, "Retrieved block missing expected text"
            
    finally:
        cleanup_files(files)
        print("\nCleaned up testing files.")

if __name__ == "__main__":
    run_test()
